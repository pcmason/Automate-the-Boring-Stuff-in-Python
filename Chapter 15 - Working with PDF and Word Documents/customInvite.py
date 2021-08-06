#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
customInvite.py
Created on Thu Aug  5 22:44:54 2021

@author: paulmason

program that opens file guests.txt and creates a Word document with each guest having their
own personal invite per page (should create a 5 page Word doc)

steps:
    1. read in the guest.txt file and get each string into a list using the readlines() method
    2. create a blank Word document and loop through adding the text and styles for each invitation
    3. save the Word document
"""
#import the docx module to create the Word document, edit the styles and save the Document
import docx
#import a way to align added paragraphs to a Word document
from docx.enum.text import WD_ALIGN_PARAGRAPH
#imoprt the pt for Font
from docx.shared import Pt

#tell the user what the program does
print('This program opens the file guests.txt and creates a Word doc with invitations for each guest.')
#step 1 - read in the guests
#open the guests.txt file in read mode
guestFile = open('guests.txt', 'r')
#use the readlines method to get each guests name stored in a list
guests = guestFile.readlines()

#step 2 - open a blank Word document and add each guests invitation to it 
#create a new Word document
doc = docx.Document()

#loop through the guests in the list
for guest in guests:
    #output to user what invitation is being created
    print('Creating invitation for %s...' % guest)
    #add the first paragraph with the invit
    paraIntro = doc.add_paragraph('It would be a pleasure to have the company of')
    #set the style of the intro paragraph to quote
    paraIntro.style = 'Quote'
    #align the intro paragraph
    paraIntro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #add the guests name to the invite give it the style MacroText
    paraGuest = doc.add_paragraph(guest, 'MacroText')
    #center the guest paragraph
    paraGuest.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #increase the font of the guest paragraph
    paraGuest.runs[0].font.size = Pt(24)
    #create the paragraph for location, give it the style of Quote
    paraLoc = doc.add_paragraph('at', 'Quote')
    #add a run for the second part of the paragraph
    paraLoc.add_run(' 11010 Memory Lane on the Evening of')
    #underline the at
    paraLoc.runs[0].underline = True
    #set the style of the second run object to quote
    paraLoc.runs[1].style = 'QuoteChar'
    #center align the location
    paraLoc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #add the data paragraph, give it the Book Title style
    paraDate = doc.add_paragraph('April 1st', 'TOC Heading')
    #align the paragraph date 
    paraDate.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #add  a paragraph for the time
    paraTime = doc.add_paragraph('at', 'Quote')
    #add the run for the rest of the time paragraph
    paraTime.add_run(" 7 o'clock")
    #underline the at 
    paraTime.runs[0].underline = True
    #set the style of the second run object to quote
    paraTime.runs[1].style = 'QuoteChar'
    #center align the time
    paraTime.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #create a page break now for the next invitation
    paraTime.runs[1].add_break(docx.enum.text.WD_BREAK.PAGE)

#step 3 - save the invitations doc    
#now save the Word document title it 'invitations.docx'
doc.save('invitations.docx')
#output to the user the program is done
print('Done! Find invitations in the cwd as \'invitations.docx\'.')